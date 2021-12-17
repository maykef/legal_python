# https://python-docx.readthedocs.io/en/latest/index.html

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn

document = Document()

style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

document.add_picture('Uni_logo.png', width=Cm(4.1))
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
last_paragraph.paragraph_format.space_before = Pt(80)
last_paragraph.paragraph_format.space_after = Pt(40)

p = document.add_paragraph()
run = p.add_run('DATED')
p.paragraph_format.left_indent = Cm(2.5)
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(60)
run.bold = True

p = document.add_paragraph()
run = p.add_run('(1) THE CHANCELLOR, MASTERS AND SCHOLARS OF')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True
run.add_break()
run = p.add_run('THE UNIVERSITY OF CAMBRIDGE')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True


p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('-and-')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('(2) [CONTRACTOR]')
run.bold = True
last_paragraph.paragraph_format.space_before = Pt(20)
last_paragraph.paragraph_format.space_after = Pt(80)

p = document.add_paragraph()
run = p.add_run('[EQUIPMENT] [IT SYSTEM] [SUPPLY] [SERVICES] AGREEMENT')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True
run.add_break()
run = p.add_run('relating to')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = False
run.add_break()
run = p.add_run('[the provision of ……….]')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True
run.add_break()
run = p.add_run('[for the Department of ………]')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = True
run.add_break()
run = p.add_run('AT')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = False
run.add_break()
run = p.add_run('THE UNIVERSITY OF CAMBRIDGE')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Contract Reference Number: ')
run.bold = False
run.add_break()
run = p.add_run('UCAM xxx/xx or DISP xxxxx')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
last_paragraph.paragraph_format.space_after = Pt(60)
run.bold = False

document.add_page_break()

#p = document.add_paragraph()
#last_paragraph.paragraph_format.space_before = Pt(20)
#last_paragraph.paragraph_format.space_after = Pt(20)

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(20)
last_paragraph.paragraph_format.space_after = Pt(20)
run = p.add_run('THIS AGREEMENT dated the                    \
                                    (the “Effective Date”) is')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('BETWEEN:')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('1.\tTHE CHANCELLOR, MASTERS AND SCHOLARS \
OF THE UNIVERSITY OF \tCAMBRIDGE of The Old Schools, \
Trinity Lane, Cambridge CB2 1TN \t("University");')
run.bold = False



p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('2.\t[THE CONTRACTOR]')
run.bold = False
run.add_break()
run = p.add_run('\tof ("The Contractor");')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False



p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run('3.\t[THE GUARANTOR]')
run.bold = False
run.add_break()
run = p.add_run('\tof ("The Guarantor");')
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('BACKGROUND')
run.bold = True

## This is a workaround as the number list cannot be reset.
## style='List Number' deleted from the add_paragraph
## runs have been added to be able to align the text.
p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('1.\tBy an Invitation dated    day of         20[xx] \
(“Invitation”) the University invited')
run.add_break()
run = p.add_run('\tthe Contractor to submit an offer\
to make a supply [of Goods [a System')
run.add_break()
run = p.add_run('\tcomprising software [and hardware]]\
 [and] Services [and maintenance services]')
run.add_break()
run = p.add_run('\t[and Works]] (“the Supply”)')
run.bold = False

## This is a workaround as the number list cannot be reset.
## style='List Number' deleted from the add_paragraph
## runs have been added to be able to align the text.
p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run('2.\tBy a Proposal dated the    day of         20[  ] \
(“Proposal”) the Contractor ')
run.add_break()
run = p.add_run('\toffered to make the Supply in accordance with the University’s conditions')
run.add_break()
run = p.add_run('\tof contract at the price and rates set out in the Contractor’s Proposal.')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('IT IS AGREED AS FOLLOWS:-')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('1.')
run.bold = False
run = p.add_run('\tCONTRACT DOCUMENTS, DEFINITIONS [AND DURATION]')
run.bold = True


p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('1.1')
run.bold = False
run = p.add_run('\tThe Contract Documents listed in the Schedule (“the Contract Documents”) ')
run.bold = False
run = p.add_run('\tconstitute the contract between the parties. In the event of a conflict between the \
\tbetween the provisions of any of the Contract Documents, the provisions of the \tContract Document listed \
higher in the list of Contract Documents shall prevail.')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('1.2\tAny schedule to this Agreement forms part of it.\
Definitions and interpretation \tprovisions in the contract terms specified in\
the Contract Documents shall apply \tto this Agreement.  Unless otherwise agreed,\
the contractual specification is \tthe specification set out in the University’s \
Invitation, except to the extent that the \tContractor’s Proposal clearly states \
that an aspect of the specification cannot be \tachieved.')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('1.3\tThis Agreement shall begin on    and end on    (“Expiry Date”)\
unless terminated \tearlier or extended in accordance with the Contract.] \
[The University reserves \tthe right to extend this Agreement beyond the \
Expiry Date for further period[s] up \tto [24] months in total. Should \
the University wish to exercise this option the \tAuthorised Officer will \
confirm the extension in writing one calendar month \tbefore the Expiry Date.\
The initial and any extended period shall constitute the \tcontract period.')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('2.')
run.bold = False
run = p.add_run('\tTHE ENGAGEMENT OF THE CONTRACTOR')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('\tThe University engages the Contractor \
and the Contractor accepts the \tengagement to provide the \
Supply on the terms and conditions set out in the \tContract \
Documents for the Contract Price.')
run.bold = False

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('3.')
run.bold = False
run = p.add_run('\tGUARANTEE OF THE CONTRACTOR´S OBLIGATIONS')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('3.1')
run.bold = False
run = p.add_run('\tThe Guarantor guarantees to the University \
the obligations of the Contractor \tunder the Contract Documents \
in the terms set out in Schedule.')
run.bold = False

p = document.add_paragraph()
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('AGREED by the parties through their authorised signatories:')
run.bold = False

p = document.add_paragraph()
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)

table = document.add_table(rows=0, cols=2, style="Table Grid")
for item in range(1, 2):
    row = table.add_row() # define row and cells separately

    # accessing row xml and setting tr height
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "3500")
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

    row_cells = row.cells
    row_cells[0].text = '\nFor and on behalf \
    of The Chancellor, Masters and Scholars \
    of the University of Cambridge\n \n Signed:           ___________________ \n \
    \n Print Name:    ___________________ \n       \nTitle:                ___________________ \n\
    \n Date:              ___________________'
    row_cells[1].text = '\nFor and on behalf of \
    [Insert name of Contractor]\n \n Signed\n:           ___________________ \n \
    \n Print Name:    ___________________ \n       \nTitle:                ___________________ \n\
    \n Date:              ___________________'

document.add_page_break()

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('THE SCHEDULE 1')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
last_paragraph.paragraph_format.space_before = Pt(10)
last_paragraph.paragraph_format.space_after = Pt(10)
run = p.add_run('(CONTRACT DOCUMENTS)')
run.bold = True

p = document.add_paragraph()
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
last_paragraph.paragraph_format.space_before = Pt(30)
last_paragraph.paragraph_format.space_after = Pt(30)
run = p.add_run('1.\tThis Form of Agreement \r')
run.bold = False
run.add_break()
run = p.add_run('2.\tThe University Conditions of Contract \r')
run.bold = False
run.add_break()
run = p.add_run('3.\tPurchase Order [and Attachment] \r')
run.bold = False
run.add_break()
run = p.add_run('4.\tUniversity’s Invitation Document \
[ref: insert details for example Invitation to \tTender dated 20[  ]] \r')
run.bold = False
run.add_break()
run = p.add_run('5.\tContractor´s Proposal Document \
[ref: insert details for example Tender \tResponse Document dated  \
                                  200[  ]  \r')
run.bold = False
run.add_break()

document.save('demo_2.docx')
